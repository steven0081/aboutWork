import docx, os
from docx import Document

file_name =  os.getcwd()+'\\' + 'A1-华新水泥大数据平台-ok.docx'
doc = docx.Document(file_name)
paragraphs = doc.paragraphs
print(len(paragraphs))
print(paragraphs[4].text)
document = Document(file_name)
tables = document.tables
print('#################################处理表格#############################')
print(len(tables))
#table = tables[0]
for table in tables:
    for i in range(0, len(table.rows)):
        print(i)
        row_cells = table.rows[i].cells
        for cell in row_cells:
            print(cell)

