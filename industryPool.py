#pip install pdfplumber
import docx
import pdfplumber
import os
import openpyxl
from docx import Document

def read_docxFile ():
    #初始化EXCEL文件相关信息
    xlsx_file = r'E:\joeCloud\Documents\工业互联网\工业互联网资源池\第二批\temp.xlsx'
    wb = openpyxl.load_workbook(xlsx_file)
    sheet = wb.get_sheet_by_name('Sheet1')
    row_no = sheet.max_row
    col_no = sheet.max_column
    #设置ＤＯＣＸ文件信息
    file_dir = r'E:\joeCloud\Documents\工业互联网\工业互联网资源池\第二批\收到材料\docx'
    file_list = os.listdir(file_dir)
    for file in file_list:
        file_name = os.path.join(file_dir, file)
        print(file_name)
        #doc = docx.Document(file_name)
        document = Document(file_name)
        tables = document.tables
        #print(len(tables))
        for tab_index in range(0, 2):
            table = tables[tab_index]
            row_count = len(table.rows)
            col_count = len(table.columns)+1
            for i in range(row_count):
                cells = table.rows[i].cells
                j = 0
                txt_new, txt_old = '', ''
                for cell in cells:
                    txt_new = cell.text
                    if txt_new != txt_old:
                        sheet.cell(row=row_no + i + 1, column=j + 1).value = txt_new
                        txt_old = txt_new
                    j = j+1
            row_no = row_no + row_count
        row_no = row_no + 2
    wb.save(xlsx_file)

def read_pdfFile ():
    # 初始化EXCEL文件相关信息
    xlsx_file = r'E:\joeCloud\Documents\工业互联网\工业互联网资源池\第二批\temp.xlsx'
    wb = openpyxl.load_workbook(xlsx_file)
    sheet = wb.get_sheet_by_name('Sheet1')
    row_no = sheet.max_row +2
    #print(row_no)
    # 设置 PDF 文件信息
    file_dir = r'E:\joeCloud\Documents\工业互联网\工业互联网资源池\第二批\收到材料\pdf'
    file_list = os.listdir(file_dir)
    for file in file_list:
        pdf_file = os.path.join(file_dir, file)
        #pdf_file = r'E:\joeCloud\Documents\工业互联网\工业互联网资源池\第二批\收到材料\云南广数科技有限公司工业互联网服务商资源池申报20190715.pdf'
        pdf = pdfplumber.open(pdf_file)
        t_count = 0
        for page in pdf.pages:
            # 获取当前页面的全部文本信息，包括表格中的文字
            # print(page.extract_text())
            for table in page.extract_tables():
                r_count = 1
                for row in table:
                    #print(r_count)
                    for i in range(len(row)):
                        if row[i] is not None:
                            sheet.cell(row=row_no+r_count, column=i+1).value = row[i]
                    #print(row)
                    r_count += 1
                row_no = row_no + r_count + 2
                t_count = t_count + 1
                #print('t_count= ', t_count)
                #print('r_count=', r_count)
                #print('row_no = ', row_no)
                print('---------- 分割线 ----------')
            if t_count > 2 :
                break
        pdf.close()
        #关闭 EXCEL 文件
        wb.save(xlsx_file)

def format_xlsFile():
    #把读取的每个单位信息转化为列表
    xlsx_file = r'E:\joeCloud\Documents\工业互联网\工业互联网资源池\第二批\temp.xlsx'
    wb = openpyxl.load_workbook(xlsx_file)
    sheet = wb.get_sheet_by_name('Sheet2')
    row_no = sheet.max_row + 1
    col_no = sheet.max_column
    for i in range(3,row_no):
        name = sheet.cell(row= i , column= 2).value
        print(name)
        s = wb.get_sheet_by_name('Sheet1')
        row = s.max_row
        for j in range(row):
            name_2 = s.cell(row= j+1 , column= 2).value
            if name == name_2:
                print(j)
                #分别取：  申报方向	单位性质	联系人	电话	主营业务收入	员工总数	成立时间   企业简介
                sheet.cell(row= i , column= 4).value =s.cell(row= j , column= 2).value
                sheet.cell(row=i, column=5).value = s.cell(row=j + 3, column=2).value
                sheet.cell(row=i, column=6).value = s.cell(row=j + 7, column=3).value
                sheet.cell(row=i, column=7).value = s.cell(row=j + 8, column=5).value
                sheet.cell(row=i, column=8).value = s.cell(row=j + 4, column=2).value
                sheet.cell(row=i, column=9).value = s.cell(row=j + 4, column=6).value
                sheet.cell(row=i, column=10).value = s.cell(row=j + 3, column=6).value
                sheet.cell(row=i, column=13).value = s.cell(row=j + 10, column=2).value
    wb.save(xlsx_file)


if __name__ == "__main__":
    #read_docxFile()
    #read_pdfFile()
    format_xlsFile()




